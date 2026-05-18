"""End-to-end test for the DOCX path. Builds a tiny .docx in a temp dir."""

from __future__ import annotations

from pathlib import Path

import pytest

docx = pytest.importorskip("docx")  # python-docx

from doc2latex.pipeline import ConvertOptions, run_conversion


def _make_sample_docx(path: Path) -> None:
    from docx import Document

    d = Document()
    d.add_heading("doc2latex demo", level=1)
    d.add_paragraph("This is a paragraph with some text.")
    d.add_heading("A subsection", level=2)
    d.add_paragraph("First bullet", style="List Bullet")
    d.add_paragraph("Second bullet", style="List Bullet")
    d.add_paragraph("Closing paragraph with $special_chars & friends%.")
    d.save(str(path))


def test_docx_end_to_end(tmp_path: Path):
    inp = tmp_path / "sample.docx"
    _make_sample_docx(inp)

    out = tmp_path / "sample.tex"
    assets = tmp_path / "assets"
    opts = ConvertOptions(input=inp, out=out, assets_dir=assets)
    result = run_conversion(opts)

    text = out.read_text(encoding="utf-8")
    assert r"\documentclass" in text
    assert r"\section{doc2latex demo}" in text
    assert r"\subsection{A subsection}" in text
    assert r"\begin{itemize}" in text
    # special chars must be escaped
    assert r"\$special\_chars" in text
    assert result.block_count > 0
