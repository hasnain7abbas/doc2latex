"""DOCX reader using python-docx.

Maps Word styles → LaTeX block kinds. Extracts embedded images from the
zipped docx package directly so we don't depend on COM/Word automation.
"""

from __future__ import annotations

import zipfile
from pathlib import Path
from typing import Iterator, List, TYPE_CHECKING

from doc2latex.blocks import Block

if TYPE_CHECKING:
    from rich.console import Console

    from doc2latex.pipeline import ConvertOptions


def _heading_level(style_name: str) -> int:
    """Return 1..6 if the style is a heading, else 0."""
    if not style_name:
        return 0
    name = style_name.strip().lower()
    if name == "title":
        return 1
    if name.startswith("heading "):
        try:
            return min(6, max(1, int(name.split()[1])))
        except ValueError:
            return 0
    return 0


def _is_list(style_name: str) -> tuple[bool, bool]:
    """Return (is_list, is_ordered)."""
    if not style_name:
        return False, False
    name = style_name.lower()
    if "list number" in name or "list paragraph" in name and "number" in name:
        return True, True
    if "list bullet" in name or "list paragraph" in name:
        return True, False
    return False, False


def _runs_to_text(paragraph) -> str:
    """Concatenate runs and apply bold/italic LaTeX markup.

    We escape special characters at assembly time (text_to_latex.escape), but
    \\textbf / \\textit cannot survive that escape. So we leave plain text here
    and let the assembler escape it — bold/italic markup is dropped in this
    simple path. (Word-level inline math is handled separately by extracting
    embedded equation images.)
    """
    return paragraph.text


class DocxReader:
    def read(
        self, opts: "ConvertOptions", console: "Console"
    ) -> Iterator[Block]:
        try:
            from docx import Document  # python-docx
        except ImportError as e:  # pragma: no cover
            raise RuntimeError(
                "python-docx is not installed — required for DOCX input."
            ) from e

        doc = Document(str(opts.input))
        assets_dir: Path = opts.assets_dir

        # ---------- extract embedded media into assets_dir ----------
        media_paths: list[Path] = []
        with zipfile.ZipFile(str(opts.input)) as zf:
            for name in zf.namelist():
                if name.startswith("word/media/"):
                    base = Path(name).name
                    out = assets_dir / base
                    out.write_bytes(zf.read(name))
                    media_paths.append(out)

        # ---------- iterate paragraphs and tables in document order ----------
        # python-docx exposes body children via doc.element.body — we iterate
        # them in source order to preserve table position.
        body = doc.element.body
        # Tag URIs in WordprocessingML.
        W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

        # Build a quick lookup from paragraph XML id → Paragraph object.
        para_by_xml = {p._element: p for p in doc.paragraphs}
        tbl_by_xml = {t._element: t for t in doc.tables}

        # Track current list streak so we can collapse consecutive list items.
        list_items: list[str] = []
        list_ordered = False

        def flush_list() -> Block | None:
            nonlocal list_items, list_ordered
            if not list_items:
                return None
            blk = Block(
                kind="list",
                ordered=list_ordered,
                items=list_items,
            )
            list_items = []
            list_ordered = False
            return blk

        for child in body.iterchildren():
            tag = child.tag
            if tag == W_NS + "p":
                p = para_by_xml.get(child)
                if p is None:
                    continue
                text = _runs_to_text(p).strip()
                style_name = p.style.name if p.style else ""
                if not text:
                    # Empty paragraph — flush any pending list.
                    flushed = flush_list()
                    if flushed:
                        yield flushed
                    continue

                lvl = _heading_level(style_name)
                if lvl:
                    flushed = flush_list()
                    if flushed:
                        yield flushed
                    yield Block(kind="heading", content=text, level=lvl)
                    continue

                is_list, ordered = _is_list(style_name)
                if is_list:
                    if list_items and ordered != list_ordered:
                        yield flush_list()  # type: ignore[misc]
                    list_ordered = ordered
                    list_items.append(text)
                    continue

                flushed = flush_list()
                if flushed:
                    yield flushed
                yield Block(kind="para", content=text)

            elif tag == W_NS + "tbl":
                flushed = flush_list()
                if flushed:
                    yield flushed
                t = tbl_by_xml.get(child)
                if t is None:
                    continue
                rows: list[list[str]] = []
                for row in t.rows:
                    rows.append([cell.text.strip() for cell in row.cells])
                yield Block(kind="table", rows=rows)

        # End-of-doc flush.
        flushed = flush_list()
        if flushed:
            yield flushed

        # ---------- emit figure blocks for media that didn't get inlined ----------
        # Simple, predictable path: list every extracted image as a figure at the end.
        # (Word's image-anchor model is complex; this avoids surprising users.)
        for p in media_paths:
            # Skip non-image artifacts (charts, themes, etc.).
            if p.suffix.lower() not in {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".emf", ".wmf"}:
                continue
            if p.suffix.lower() in {".emf", ".wmf"}:
                # EMF/WMF won't embed in pdflatex without conversion; skip with note.
                if opts.verbose:
                    console.print(f"[yellow]skipping vector media:[/yellow] {p.name}")
                continue
            yield Block(kind="figure", content=str(p))
